import pandas as pd
import os
import calendar
from dateutil.relativedelta import relativedelta


from num2words import num2words
from openpyxl import load_workbook
from datetime import date,datetime,timedelta
from docx import Document

BILLS_FOLDER_PATH = os.path.join("bill docs")


def generate_data(rar_template_path,rar_no):
    data = pd.read_excel(rar_template_path,skiprows=10,header=1,skipfooter=7).fillna("")
    data.columns = ["Sl.No","Description","Qty","Unit","Rate","Amount","Qty_Executed","Qty_Passed","Amount_Executed","Qty_Last","Amount_Last","Qty_present","Amount_present"]

    data.loc[:8,"Qty_Executed"] = data.loc[:8,"Qty"] * rar_no / 8
    data["Qty_Passed"] = data["Qty_Executed"]
    data.loc[:8,"Amount_Executed"] = data.loc[:8,"Qty_Passed"] * data.loc[:8,"Rate"]

    data.loc[:8,"Qty_Last"] = data.loc[:8,"Qty"] * (rar_no-1) / 8
    data.loc[:8,"Amount_Last"] = data.loc[:8,"Qty_Last"] * data.loc[:8,"Rate"]

    data.loc[:8,"Qty_present"] = data.loc[:8,"Qty"] * (1) / 8
    data.loc[:8,"Amount_present"] = data.loc[:8,"Qty_present"] * data.loc[:8,"Rate"]

    data.loc[:8,["Sl.No"]] = list(range(1,10))

    amount_cols = [col for col in data.columns if col.find("Amount") != -1]
    data.loc[9,amount_cols] = data.loc[:8,amount_cols].sum()
    data.loc[10,amount_cols] = data.loc[9,amount_cols]*0.18
    data.loc[11,amount_cols] = data.loc[9:10,amount_cols].sum()
    return data


def get_rar_dates(start_date =  date(2022,8,10),billing_freq=3):
    dates = []

    for i in range(8):

        end_date = start_date + pd.DateOffset(months=billing_freq,days=-1)
        dates.append([start_date,end_date])
        start_date = end_date + pd.DateOffset(months=0,days=1)

    dates = pd.DataFrame(dates,columns=["start_date","end_date"],index=range(1,9))
    return dates


def generate_abstract(data,rar_template_path,rar_no,penalty,target_path):
    
    # Update with your file's path
    workbook = load_workbook(rar_template_path)
    sheet = workbook['RAR']  # Update with your sheet's name

    start_row = 13
    start_column = 3  # Assuming you start from column A

    for index, row in data.iterrows():
        for col_num, value in enumerate(row[2:], start=start_column):
            cell = sheet.cell(row=start_row + index, column=col_num)
            cell.value = value

    sheet.unmerge_cells('A2:M2') 
    # Change the value of the top-left cell
    sheet['A2'] = f'{rar_no}Th RAR'
    # Merge the cells again if needed
    sheet.merge_cells('A2:M2')

    sheet["M7"] = date.today().strftime("%d-%m-%Y")
    
    dates = get_rar_dates()
    st_dt,end_dt = dates.loc[rar_no]["start_date"].strftime("%d-%m-%Y"), dates.loc[rar_no]["end_date"].strftime("%d-%m-%Y")
    text = f"(period from Dt.{st_dt} to Dt.{end_dt})"

    sheet.unmerge_cells('I9:M9') 
    # Change the value of the top-left cell
    sheet['I9'] = text
    # Merge the cells again if needed
    sheet.merge_cells('I9:M9')

    payment_amount = data.loc[11,"Amount_present"]  - (data.loc[11,"Amount_present"] *0.05) - penalty
    # Update with your file's path
    sheet_2 = workbook['Payment Recommendation']  # Update with your sheet's name
    sheet_2["C7"] = penalty
    sheet_2["B9"] = sheet_2["B9"].value.replace("?",str(rar_no))
    sheet_2["B12"] = sheet_2["B12"].value.replace("bill_amount",str(payment_amount)).replace("?",str(rar_no))
    sheet_2["B11"] = "Rupees " + num2words(payment_amount,lang='en_IN') + "paisa only"
    workbook.save(filename=target_path)
    
    return "Bill Abstract & Payment Recommendation File Generated Successfully!"


def generate_recommendation(data,penalty,rar_template_path,rar_no):
    payment_amount = data.loc[11,"Amount_present"]  - (data.loc[11,"Amount_present"] *0.05) - penalty
    # Update with your file's path
    workbook = load_workbook(rar_template_path)
    sheet_2 = workbook['Payment Recommendation']  # Update with your sheet's name
    sheet_2["C7"] = penalty
    sheet_2["B9"] = sheet_2["B9"].value.replace("?",str(rar_no))
    sheet_2["B12"] = sheet_2["B12"].value.replace("bill_amount",str(payment_amount)).replace("?",str(rar_no))
    sheet_2["B11"] = "Rupees " + num2words(payment_amount,lang='en_IN') + "paisa only"
    workbook.save(filename=f"RAR_{rar_no}/abstract.xlsx")
    return "Bill Recommendation File Generated Successfully!"


def generate_checklist(checklist_template_path,rar_no):
    dates = get_rar_dates()
    st_dt,end_dt = dates.loc[rar_no]["start_date"].strftime("%d-%m-%Y"), dates.loc[rar_no]["end_date"].strftime("%d-%m-%Y")

    workbook = load_workbook(checklist_template_path)
    sheet = workbook[workbook.sheetnames[0]]
    
    sheet["B5"] = sheet["B5"].value.replace("?",str(rar_no)).replace("date1",st_dt).replace("date2",end_dt)
    workbook.save(filename=f"RAR_{rar_no}/checklist.xlsx")
    return "Checklist file generated successfully "


def generate_undertaking(undertaking_template_path,penalty,rar_no,invoice_date,GE_NO,RR_NO,GE_date):
    
    data = generate_data(rar_template_path,rar_no)
    invoice_amount = data.iloc[-1,-1]
    ret_money = round(invoice_amount *0.05,2)
    rec_amount =  invoice_amount - penalty - ret_money
    
    workbook = load_workbook(undertaking_template_path)
    sheet = workbook[workbook.sheetnames[0]]
    sheet["C5"] = sheet["C5"].value.replace("?",str(rar_no))
    sheet["C6"] = sheet["C6"].value.replace("GE_NO",str(GE_NO)).replace("date",GE_date)
    sheet["C7"] = sheet["C7"].value.replace("RR_NO",RR_NO).replace("date",GE_date)
    sheet["C10"] = sheet["C10"].value.replace("date",invoice_date)
    sheet["C13"] = sheet["C13"].value.replace("Penalty",str(penalty)).replace("total_deduction",str(penalty+ret_money))
    sheet["C14"] = sheet["C14"].value.replace("recom_amount",str(rec_amount))
    workbook.save(filename=f"RAR_{rar_no}/undertaking.xlsx")
    return "Undertaking file generated successfully"


def generate_acceptance_report(acceptance_report_template_path,rar_no,penalty):
    dates = get_rar_dates()
    st_dt,end_dt = dates.loc[rar_no]["start_date"].strftime("%d-%m-%Y"), dates.loc[rar_no]["end_date"].strftime("%d-%m-%Y")
#     print(st_dt,end_dt)
    # Load an existing document
    doc = Document(acceptance_report_template_path)
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("?", str(rar_no))
    doc.paragraphs[6].text = doc.paragraphs[6].text.replace("?", str(rar_no)).replace("date1",st_dt).replace("date2",end_dt)
    doc.paragraphs[8].text = doc.paragraphs[8].text.replace("penalty_amt",str(penalty))
    doc.save(f"RAR_{rar_no}/acceptance_report.docx")
    return "Acceptance Report File Generated Successfully!"


def generate_attendance_data(selected_month, year=2024, employees=None):

    if employees is None:
        employees = []  # Default empty list
    
    # Dictionary mapping month names to their corresponding numbers
    month_number = {name: num for num, name in enumerate(calendar.month_name) if num}
    
    num_days = calendar.monthrange(year=year, month=selected_month)[1]
    attendance_data = []
    
    attendance_data = {"dates":[f"{day:02d}/{selected_month:02d}" for day in range(1, num_days + 1)]}
    
    values = {emp:[ 1 if calendar.weekday(year, selected_month, day) != 6 else 0 for day in range(1, num_days + 1)]  for emp in employees}
    days = {"sunday":[ "No Sunday" if calendar.weekday(year, selected_month, day) != 6 else "sunday" for day in range(1, num_days + 1)]}
    attendance_data.update(values)
    attendance_data.update(days)
    # combined_data = {}
    # for key in attendance_data.keys():
    #     combined_data[key] = zip(attendance_data[key], attendance_data['sunday'])
    return attendance_data

def map_values(value):
    if value == "PP":
        return 1
    elif value in ["AP","PP"]:
        return 0.5
    else:
        return 0
    
def get_punching_values(filename):
    punch_data = pd.read_csv(filename)
    punch_df = punch_data.map(map_values)
    values = punch_df[punch_df.columns[1:-1]].values.tolist() 
    return values

def generate_attendance_data_df(employees,month,year,contract_no,values=None):
    dates = pd.date_range(start=f"{year}-{month}-01", end=pd.Period(f"{year}-{month}").end_time, freq='D')
    attendance_data = pd.DataFrame(columns=employees,index=dates).reset_index().rename(columns={"index":"date"})
    attendance_data["sunday"] = attendance_data["date"].dt.day_name().apply(lambda x:"sunday" if x == "Sunday" else "weekday")
    attendance_data["date_format"] = attendance_data["date"].dt.strftime("%d/%m")
    attendance_folder = os.path.join(BILLS_FOLDER_PATH,contract_no,"attendance")
    punch_data_filename = os.path.join(attendance_folder,"updated_punch_data.csv")

    if values == None:
        attendance_folder = os.path.join(BILLS_FOLDER_PATH, contract_no, "attendance")
        file_path = os.path.join(attendance_folder,f"{month}-{year}.csv")
        #if os.path.exists(file_path):
        try:
            data = pd.read_csv(file_path)
            #values = [list(data[col].values) for col in data.columns[1:-2]]
            values = [list(data[col].values) for col in employees]

        except:
            values = get_punching_values(punch_data_filename)

            #values = [[1 if day != "sunday" else 0 for day in attendance_data["sunday"]] for _ in range(len(employees))]

    present_days = ["total_days"]

    for index,col in enumerate(employees):
        attendance_data[col] = values[index]
        present_days.append(sum(values[index]))

    present_days.extend(["weekday","DP"])
    attendance_data.loc["total"] = present_days
    
    attendance_dict = {col:list(attendance_data[col].values) for col in attendance_data.columns[1:] }

    return attendance_dict

def get_prev_months():
    # Get the current date
    current_date = datetime.now()
    # List to store the last 3 months
    last_3_months = []
    # Loop to generate the last 3 months
    for i in range(3):
        # Subtract i months from the current date
        month = current_date.month - i
        year = current_date.year
        if month <= 0:
            # If the month goes below 1, adjust the year accordingly
            month += 12
            year -= 1
        last_3_months.append((year, month))

    return {calendar.month_name[month]: f"{month}-{year}" for year,month in last_3_months}


def get_month_name(month_no):
    return calendar.month_name[month_no]

def split_list(lst, num_sublists):
    # Calculate the length of each sublist
    sublist_length = len(lst) // num_sublists
    remainder = len(lst) % num_sublists
    # Initialize the starting index
    start = 0
    sublists = []
    # Create sublists
    for i in range(num_sublists):
        # Calculate the ending index for this sublist
        end = start + sublist_length + (1 if i < remainder else 0)
        # Append the sublist to the result
        sublists.append(lst[start:end])
        # Update the starting index for the next sublist
        start = end
    return sublists


def attendance_to_csv(month,year,contract_no,df):
    attendance_folder = os.path.join(BILLS_FOLDER_PATH, contract_no, "attendance")
    os.makedirs(attendance_folder,exist_ok=True)
    file_path = os.path.join(attendance_folder,f"{month}-{year}.csv")
    df.to_csv(file_path)

    return "Attendance Saved Successfully!"

